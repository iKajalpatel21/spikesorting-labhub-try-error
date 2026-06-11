import io
import os
import re
import tempfile
from html.parser import HTMLParser

from django.http import HttpResponse
from rest_framework import status
from rest_framework.decorators import api_view, parser_classes, permission_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .models import LabNote, NoteAuditLog
from .serializers import LabNoteSerializer


def _log(note, user, action):
    NoteAuditLog.objects.create(note=note, user=user, action=action)


# ---------------------------------------------------------------------------
# Summarisation — lazy-loaded, cached for the process lifetime
# ---------------------------------------------------------------------------

_summarizer_cache: dict = {"tokenizer": None, "model": None, "loaded": False}

_SUMMARIZER_MODEL = "sshleifer/distilbart-cnn-12-6"


def _get_summarizer():
    if not _summarizer_cache["loaded"]:
        _summarizer_cache["loaded"] = True
        try:
            from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
            _summarizer_cache["tokenizer"] = AutoTokenizer.from_pretrained(_SUMMARIZER_MODEL)
            _summarizer_cache["model"] = AutoModelForSeq2SeqLM.from_pretrained(_SUMMARIZER_MODEL)
            _summarizer_cache["model"].eval()
        except Exception:
            pass  # model stays None; summarisation gracefully disabled
    return _summarizer_cache["tokenizer"], _summarizer_cache["model"]


def _summarize_to_bullets(text: str) -> str | None:
    """
    Returns an HTML <ul> of bullet-point sentences summarising *text*,
    or None when the text is too short or the model is unavailable.
    """
    words = text.split()
    if len(words) < 30:
        return None

    tokenizer, model = _get_summarizer()
    if tokenizer is None or model is None:
        return None

    # BART's max input is ~1 024 tokens; cap at ~650 words to stay safe
    if len(words) > 650:
        text = " ".join(words[:650])

    try:
        import torch
        inputs = tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
        with torch.no_grad():
            ids = model.generate(
                inputs["input_ids"],
                max_length=150, min_length=30,
                num_beams=4, early_stopping=True,
            )
        summary_text = tokenizer.decode(ids[0], skip_special_tokens=True).strip()
    except Exception:
        return None

    # Split on sentence boundaries → one bullet per sentence
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', summary_text) if s.strip()]
    if not sentences:
        return None

    items = "".join(f"<li>{s}</li>" for s in sentences)
    return f"<ul>{items}</ul>"


# ---------------------------------------------------------------------------
# HTML → reportlab paragraph converter
# ---------------------------------------------------------------------------

class _HtmlToRL(HTMLParser):
    _BLOCK = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr", "td"}
    _OPEN  = {"b": "<b>", "strong": "<b>", "i": "<i>", "em": "<i>", "u": "<u>"}
    _CLOSE = {"b": "</b>", "strong": "</b>", "i": "</i>", "em": "</i>", "u": "</u>"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._paragraphs = []
        self._buf = ""

    def _flush(self):
        self._paragraphs.append(self._buf.strip())
        self._buf = ""

    def handle_starttag(self, tag, attrs):
        t = tag.lower()
        if t in self._BLOCK or t == "br":
            self._flush()
        elif t in self._OPEN:
            self._buf += self._OPEN[t]

    def handle_endtag(self, tag):
        t = tag.lower()
        if t in self._BLOCK:
            self._flush()
        elif t in self._CLOSE:
            self._buf += self._CLOSE[t]

    def handle_data(self, data):
        self._buf += data.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def paragraphs(self):
        self._flush()
        return self._paragraphs


def _compliance_lines(note) -> list[str]:
    """Return the standard audit metadata lines shared across all export formats."""
    edit_count = note.audit_logs.filter(action=NoteAuditLog.ACTION_EDITED).count()
    return [
        f"Note ID:       #{note.id}",
        f"Author:        {note.created_by.username}",
        f"Created:       {note.created_at.strftime('%B %d, %Y  %I:%M %p')} UTC",
        f"Last modified: {note.updated_at.strftime('%B %d, %Y  %I:%M %p')} UTC",
        f"Revisions:     {edit_count}",
    ]


def _build_pdf(note) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    title_style      = ParagraphStyle("NoteTitle", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=2,  textColor=colors.HexColor("#1e1e1e"))
    compliance_style = ParagraphStyle("Compliance", parent=styles["Normal"],  fontSize=8,  leading=12, spaceAfter=2,  textColor=colors.HexColor("#888888"))
    body_style       = ParagraphStyle("NoteBody",  parent=styles["Normal"],   fontSize=11, leading=17, spaceAfter=6,  textColor=colors.HexColor("#1e1e1e"))

    parser = _HtmlToRL()
    parser.feed(note.content)

    story = [Paragraph(note.title, title_style)]

    # Compliance block
    for line in _compliance_lines(note):
        story.append(Paragraph(line, compliance_style))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=12))

    for text in parser.paragraphs():
        if text:
            try:
                story.append(Paragraph(text, body_style))
            except Exception:
                story.append(Paragraph(re.sub(r"<[^>]+>", "", text), body_style))
        else:
            story.append(Spacer(1, 8))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch)
    doc.build(story)
    buf.seek(0)
    return buf.read()


def _build_docx(note) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from html.parser import HTMLParser as _HP

    doc = Document()

    # Title
    title_par = doc.add_paragraph()
    run = title_par.add_run(note.title)
    run.bold = True
    run.font.size = Pt(18)

    # Compliance block
    for line in _compliance_lines(note):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Divider (paragraph of underscores)
    div = doc.add_paragraph()
    div_run = div.add_run("─" * 60)
    div_run.font.size = Pt(8)
    div_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()  # spacer

    # Parse HTML content into docx runs
    class _DocxParser(_HP):
        _OPEN_MAP  = {"b": "bold", "strong": "bold", "i": "italic", "em": "italic", "u": "underline"}
        _CLOSE_MAP = {"b": "bold", "strong": "bold", "i": "italic", "em": "italic", "u": "underline"}
        BLOCK = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li"}

        def __init__(self, document):
            super().__init__(convert_charrefs=True)
            self._doc = document
            self._par = document.add_paragraph()
            self._fmt = {"bold": False, "italic": False, "underline": False}

        def _new_para(self):
            self._par = self._doc.add_paragraph()

        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t in self.BLOCK or t == "br":
                self._new_para()
            elif t in self._OPEN_MAP:
                self._fmt[self._OPEN_MAP[t]] = True

        def handle_endtag(self, tag):
            t = tag.lower()
            if t in self.BLOCK:
                self._new_para()
            elif t in self._CLOSE_MAP:
                self._fmt[self._CLOSE_MAP[t]] = False

        def handle_data(self, data):
            if not data.strip() and not data:
                return
            run = self._par.add_run(data)
            run.bold      = self._fmt["bold"]
            run.italic    = self._fmt["italic"]
            run.underline = self._fmt["underline"]
            run.font.size = Pt(11)

    _DocxParser(doc).feed(note.content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Speech-to-text transcription
# ---------------------------------------------------------------------------

@api_view(["POST"])
@parser_classes([MultiPartParser])
@permission_classes([IsAuthenticated])
def transcribe_audio(request):
    """
    POST multipart/form-data with field 'audio' (webm/ogg/wav/mp4).
    Returns {"text": "<transcript>"}.
    Requires openai-whisper and ffmpeg to be installed on the server.
    """
    if "audio" not in request.FILES:
        return Response({"error": "No audio file provided (field name: audio)"}, status=status.HTTP_400_BAD_REQUEST)

    audio_file = request.FILES["audio"]

    # Determine extension from content type for ffmpeg
    content_type = audio_file.content_type or ""
    ext_map = {
        "audio/webm":  ".webm",
        "audio/ogg":   ".ogg",
        "audio/wav":   ".wav",
        "audio/mp4":   ".mp4",
        "audio/mpeg":  ".mp3",
        "video/webm":  ".webm",
    }
    ext = ext_map.get(content_type, ".webm")

    try:
        import whisper
    except ImportError:
        return Response(
            {"error": "openai-whisper is not installed on this server. Run: pip install openai-whisper"},
            status=status.HTTP_501_NOT_IMPLEMENTED,
        )

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        for chunk in audio_file.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    file_size = os.path.getsize(tmp_path)
    if file_size < 1024:
        os.unlink(tmp_path)
        return Response({"text": "", "debug": f"audio too small: {file_size} bytes"})

    try:
        model = whisper.load_model("base")
        result = model.transcribe(
            tmp_path,
            language="en",           # skip language detection — avoids silent failures
            fp16=False,              # CPU-safe
            no_speech_threshold=0.4, # default 0.6 — lower = less likely to suppress output
            condition_on_previous_text=False,
        )
        text = result.get("text", "").strip()
        summary = _summarize_to_bullets(text) if text else None
        return Response({"text": text, "summary": summary, "debug": f"file={file_size}B ct={content_type}"})
    except Exception as e:
        return Response({"error": f"Transcription failed: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# CRUD endpoints
# ---------------------------------------------------------------------------

@api_view(["GET", "POST"])
@permission_classes([IsAuthenticated])
def notes_list(request):
    if request.method == "GET":
        notes = LabNote.objects.filter(created_by=request.user)
        return Response(LabNoteSerializer(notes, many=True).data)

    serializer = LabNoteSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    note = serializer.save(created_by=request.user)
    _log(note, request.user, NoteAuditLog.ACTION_CREATED)
    return Response(serializer.data, status=status.HTTP_201_CREATED)


@api_view(["GET", "PUT", "DELETE"])
@permission_classes([IsAuthenticated])
def note_detail(request, note_id):
    try:
        note = LabNote.objects.get(pk=note_id, created_by=request.user)
    except LabNote.DoesNotExist:
        return Response({"error": "Not found"}, status=status.HTTP_404_NOT_FOUND)

    if request.method == "GET":
        return Response(LabNoteSerializer(note).data)

    if request.method == "PUT":
        serializer = LabNoteSerializer(note, data=request.data, partial=True)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        serializer.save()
        _log(note, request.user, NoteAuditLog.ACTION_EDITED)
        return Response(serializer.data)

    note.delete()
    return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def note_history(request, note_id):
    try:
        note = LabNote.objects.get(pk=note_id, created_by=request.user)
    except LabNote.DoesNotExist:
        return Response({"error": "Not found"}, status=status.HTTP_404_NOT_FOUND)

    logs = note.audit_logs.select_related("user").order_by("timestamp")
    data = [
        {
            "action":    entry.action,
            "user":      entry.user.username if entry.user else "(deleted)",
            "timestamp": entry.timestamp.strftime("%b %d, %Y  %I:%M %p") + " UTC",
        }
        for entry in logs
    ]
    return Response(data)


# ---------------------------------------------------------------------------
# Export endpoints
# ---------------------------------------------------------------------------

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def download_pdf(request, note_id):
    try:
        note = LabNote.objects.get(pk=note_id, created_by=request.user)
    except LabNote.DoesNotExist:
        return Response({"error": "Not found"}, status=status.HTTP_404_NOT_FOUND)

    try:
        pdf_bytes = _build_pdf(note)
    except Exception as e:
        return Response({"error": f"PDF generation failed: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    safe_title = "".join(c for c in note.title if c.isalnum() or c in " _-")[:40].strip() or "note"
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="lab_note_{note.id}_{safe_title}.pdf"'
    return response


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def download_docx(request, note_id):
    try:
        note = LabNote.objects.get(pk=note_id, created_by=request.user)
    except LabNote.DoesNotExist:
        return Response({"error": "Not found"}, status=status.HTTP_404_NOT_FOUND)

    try:
        docx_bytes = _build_docx(note)
    except Exception as e:
        return Response({"error": f"DOCX generation failed: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    safe_title = "".join(c for c in note.title if c.isalnum() or c in " _-")[:40].strip() or "note"
    response = HttpResponse(docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="lab_note_{note.id}_{safe_title}.docx"'
    return response
